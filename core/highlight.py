"""
政策文档高亮模块
在政策原文基础上高亮华测相关的段落，输出为Word文档
"""
import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, 'policy_document_word')


def find_analysis_result(doc_path):
    """查找对应的分析结果文件"""
    result_dir = os.path.join(BASE_DIR, 'analyze_result')
    doc_name = os.path.basename(doc_path)
    doc_title = os.path.splitext(doc_name)[0]

    # 获取原始文档的相对路径
    doc_rel_dir = os.path.dirname(doc_path)
    if doc_rel_dir and doc_rel_dir != '.':
        # 先在对应的子目录查找
        target_dir = os.path.join(result_dir, doc_rel_dir)
        if os.path.exists(target_dir):
            for root, dirs, files in os.walk(target_dir):
                for f in files:
                    if f.endswith('.md') and '_分析结果_' in f:
                        for kw in [doc_title] + re.findall(r'《([^》]+)》', doc_title):
                            if kw and len(kw) >= 4 and kw in f:
                                return os.path.join(root, f)

    # 在根目录查找
    for root, dirs, files in os.walk(result_dir):
        for f in files:
            if f.endswith('.md') and '_分析结果_' in f:
                for kw in [doc_title] + re.findall(r'《([^》]+)》', doc_title):
                    if kw and len(kw) >= 4 and kw in f:
                        return os.path.join(root, f)
    return None


def parse_analysis(content):
    """解析分析结果，返回待高亮的段落列表"""
    paragraphs = []
    current_text = None
    current_score = None
    keywords = []
    current_business = None

    for line in content.split('\n'):
        line = line.strip()
        # 段落内容行
        if line.startswith('> ') and not line.startswith('> -'):
            current_text = line[2:].strip()
            keywords = []
        # 评分和关键词行
        elif line.startswith('> -') and current_text:
            match = re.search(r'评分：(\d+)/100', line)
            if match:
                current_score = int(match.group(1))
                match_kw = re.search(r'关键词：(.+)', line)
                if match_kw:
                    keywords = [k.strip() for k in match_kw.group(1).split('、')]
                match_biz = re.search(r'语义关联：(.+)', line)
                if match_biz:
                    current_business = match_biz.group(1).strip()
                if current_text and current_score is not None:
                    paragraphs.append({
                        'text': current_text,
                        'score': current_score,
                        'keywords': keywords,
                        'business': current_business
                    })
            current_text = None
            current_score = None
            keywords = []
            current_business = None
    return paragraphs


def load_original_doc(doc_path):
    """加载原始政策文档内容"""
    if not os.path.exists(doc_path):
        return None
    with open(doc_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_original_doc(content):
    """解析原始文档，返回段落列表(保留格式)"""
    paragraphs = []
    lines = content.split('\n')

    current_para = []
    in_code_block = False

    for line in lines:
        # 检测代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if current_para:
                    paragraphs.append(('\n'.join(current_para), 'code'))
                    current_para = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            current_para.append(line)
            continue

        # 空行表示段落结束
        if not line.strip():
            if current_para:
                text = '\n'.join(current_para)
                # 判断段落类型
                if text.startswith('# '):
                    ptype = 'title'
                elif text.startswith('## ') or text.startswith('**') and '：' in text:
                    ptype = 'subtitle'
                elif text.startswith('> '):
                    ptype = 'quote'
                else:
                    ptype = 'normal'
                paragraphs.append((text, ptype))
                current_para = []
            continue

        current_para.append(line)

    # 处理最后一个段落
    if current_para:
        text = '\n'.join(current_para)
        if text.startswith('# '):
            ptype = 'title'
        elif text.startswith('## ') or text.startswith('**') and '：' in text:
            ptype = 'subtitle'
        elif text.startswith('> '):
            ptype = 'quote'
        else:
            ptype = 'normal'
        paragraphs.append((text, ptype))

    return paragraphs


def find_matching_paragraph(text, highlight_list):
    """在原始段落中找到与分析结果匹配的段落"""
    text_clean = re.sub(r'\s+', ' ', text.strip())

    for idx, hl in enumerate(highlight_list):
        hl_text = hl['text']
        # 精确匹配
        if hl_text == text_clean:
            return idx
        # 包含匹配
        if hl_text in text_clean or text_clean in hl_text:
            return idx
        # 关键词匹配 (至少匹配2个关键词)
        if hl.get('keywords'):
            matches = sum(1 for kw in hl['keywords'] if kw in text_clean)
            if matches >= 2:
                return idx
    return -1


def highlight_color(score):
    """根据评分返回颜色"""
    if score >= 80:
        return RGBColor(0, 128, 0)  # 绿色
    elif score >= 60:
        return RGBColor(200, 150, 0)  # 黄色/金色
    else:
        return RGBColor(200, 0, 0)  # 红色


def create_highlighted_doc(doc_title, original_paragraphs, highlight_list, output_path):
    """创建高亮Word文档"""
    if not WORD_AVAILABLE:
        return False

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(doc_title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 匹配并高亮段落
    matched_indices = set()
    for idx, (text, ptype) in enumerate(original_paragraphs):
        # 检查是否匹配任何高亮段落
        match_idx = find_matching_paragraph(text, highlight_list)
        if match_idx >= 0:
            matched_indices.add(match_idx)
            score = highlight_list[match_idx]['score']
            color = highlight_color(score)

            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.color.rgb = color
        else:
            # 未匹配的段落正常显示
            para = doc.add_paragraph()
            run = para.add_run(text)

    # 添加高亮说明
    if matched_indices:
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("【高亮说明】")
        note_run.font.bold = True

        for idx in sorted(matched_indices):
            hl = highlight_list[idx]
            color = highlight_color(hl['score'])
            color_name = '绿色' if hl['score'] >= 80 else ('黄色' if hl['score'] >= 60 else '红色')

            note = doc.add_paragraph()
            run = note.add_run(f"• {color_name}高亮 (评分{hl['score']}/100)")
            run.font.color.rgb = color

    doc.save(output_path)
    return True


def save_analysis_result_word(content, title, output_path):
    """保存分析结果为Word文档"""
    if not WORD_AVAILABLE:
        return False

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(content)

    doc.save(output_path)
    return True


def highlight_doc(doc_path, verbose=True):
    """高亮单个文档，生成高亮Word文档"""
    if not WORD_AVAILABLE:
        return False, None, "python-docx 未安装"

    result_path = find_analysis_result(doc_path)
    if not result_path:
        if verbose:
            print(f"未找到分析结果: {doc_path}")
        return False, None, "未找到分析结果"

    # 读取分析结果并提取分数
    with open(result_path, 'r', encoding='utf-8') as f:
        analysis_content = f.read()
    score_match = re.search(r'> \*\*总分\*\*[：:]\s*([\d.]+)', analysis_content)
    score_str = score_match.group(1) if score_match else "0.0"

    doc_name = os.path.basename(doc_path)
    doc_title = os.path.splitext(doc_name)[0]

    # 获取原始文档的相对路径，保留子文件夹结构
    doc_rel_dir = os.path.dirname(doc_path)
    # 如果是绝对路径，提取相对于 policy_document 的子目录名称
    if os.path.isabs(doc_rel_dir):
        # 从绝对路径中提取子目录名称
        # 例如：C:\...\AI Skill分析\policy_document\20260203_xxx\ -> 20260203_xxx
        base_path = os.path.join(BASE_DIR, 'policy_document')
        if doc_rel_dir.startswith(base_path):
            doc_rel_dir = doc_rel_dir[len(base_path):].strip('/\\')
    else:
        # 相对路径：去掉 policy_document/ 前缀
        if doc_rel_dir.startswith('policy_document/'):
            doc_rel_dir = doc_rel_dir[len('policy_document/'):]
        elif doc_rel_dir.startswith('policy_document\\'):
            doc_rel_dir = doc_rel_dir[len('policy_document\\'):]

    if doc_rel_dir and doc_rel_dir != '.':
        target_dir = os.path.join(OUTPUT_DIR, doc_rel_dir)
    else:
        target_dir = OUTPUT_DIR

    os.makedirs(target_dir, exist_ok=True)
    output_path = os.path.join(target_dir, f"{doc_title}_高亮文档_{score_str}.docx")

    try:
        # 加载原始文档
        original_content = load_original_doc(doc_path)
        if not original_content:
            return False, None, "原始文档不存在"

        # 解析原始文档
        original_paragraphs = parse_original_doc(original_content)

        # 加载并解析分析结果
        with open(result_path, 'r', encoding='utf-8') as f:
            analysis_content = f.read()
        highlight_list = parse_analysis(analysis_content)

        if not highlight_list:
            # 没有高亮段落，直接保存原始文档
            success = save_analysis_result_word(analysis_content, doc_title, output_path)
            return success, output_path, "无高亮段落"

        # 创建高亮文档
        success = create_highlighted_doc(doc_title, original_paragraphs, highlight_list, output_path)
        if success:
            return True, output_path, f"成功 (高亮{len(highlight_list)}段)"
        return False, None, "保存失败"
    except Exception as e:
        return False, None, str(e)


def convert_analysis_to_word(doc_path, verbose=True):
    """将分析结果转换为Word文档"""
    if not WORD_AVAILABLE:
        return False, None, "python-docx 未安装"

    result_path = find_analysis_result(doc_path)
    if not result_path:
        if verbose:
            print(f"未找到分析结果: {doc_path}")
        return False, None, "未找到分析结果"

    # 读取分析结果并提取分数
    with open(result_path, 'r', encoding='utf-8') as f:
        analysis_content = f.read()
    score_match = re.search(r'> \*\*总分\*\*[：:]\s*([\d.]+)', analysis_content)
    score_str = score_match.group(1) if score_match else "0.0"

    doc_name = os.path.basename(doc_path)
    doc_title = os.path.splitext(doc_name)[0]

    # 获取原始文档的相对路径，保留子文件夹结构
    doc_rel_dir = os.path.dirname(doc_path)
    # 如果是绝对路径，提取相对于 policy_document 的子目录名称
    if os.path.isabs(doc_rel_dir):
        # 从绝对路径中提取子目录名称
        # 例如：C:\...\AI Skill分析\policy_document\20260203_xxx\ -> 20260203_xxx
        base_path = os.path.join(BASE_DIR, 'policy_document')
        if doc_rel_dir.startswith(base_path):
            doc_rel_dir = doc_rel_dir[len(base_path):].strip('/\\')
    else:
        # 相对路径：去掉 policy_document/ 前缀
        if doc_rel_dir.startswith('policy_document/'):
            doc_rel_dir = doc_rel_dir[len('policy_document/'):]
        elif doc_rel_dir.startswith('policy_document\\'):
            doc_rel_dir = doc_rel_dir[len('policy_document\\'):]

    if doc_rel_dir and doc_rel_dir != '.':
        target_dir = os.path.join(OUTPUT_DIR, doc_rel_dir)
    else:
        target_dir = OUTPUT_DIR

    os.makedirs(target_dir, exist_ok=True)
    output_path = os.path.join(target_dir, f"{doc_title}_分析结果_{score_str}.docx")

    try:
        with open(result_path, 'r', encoding='utf-8') as f:
            content = f.read()

        success = save_analysis_result_word(content, doc_title, output_path)
        if success:
            return True, output_path, "成功"
        return False, None, "保存失败"
    except Exception as e:
        return False, None, str(e)
