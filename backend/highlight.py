"""
政策文档高亮模块
在政策原文中高亮华测相关的段落，输出为Word文档

匹配策略：完全匹配
- 分析结果文本与原文完全一致才高亮
- 不进行半句、模糊、关键词等复杂匹配
"""
import os
import sys
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，将无法生成Word文档")

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, 'policy_document_word')


def find_analysis_result(doc_title):
    """查找对应的分析结果文件"""
    result_dir = os.path.join(BASE_DIR, 'analyze_result')

    for root, dirs, files in os.walk(result_dir):
        for f in files:
            if f.endswith('.md') and doc_title in f and '_分析结果_' in f:
                return os.path.join(root, f)
    return None


def parse_paragraphs(analysis_path):
    """
    从分析结果中提取原文段落、评分、关键词
    返回格式：[{'text': str, 'score': int, 'keywords': list}, ...]
    """
    paragraphs = []

    with open(analysis_path, 'r', encoding='utf-8') as f:
        content = f.read()

    if '相关段落' not in content:
        return paragraphs

    lines = content.split('\n')
    current_text = None
    current_score = None
    current_keywords = []

    for line in lines:
        line = line.strip()

        # 原文行
        if line.startswith('> ') and not line.startswith('> -'):
            current_text = line[2:].strip()
            current_keywords = []

        # 评分行
        elif line.startswith('> -') and current_text:
            score_match = re.search(r'评分[：:]\s*(\d+)', line)
            kw_match = re.search(r'关键词[：:]\s*(.+)', line)

            if score_match:
                current_score = int(score_match.group(1))
                if kw_match:
                    # 提取关键词列表
                    kw_str = kw_match.group(1).strip()
                    current_keywords = [k.strip() for k in kw_str.split('、') if k.strip()]

                paragraphs.append({
                    'text': current_text,
                    'score': current_score,
                    'keywords': current_keywords
                })

            current_text = None
            current_score = None
            current_keywords = []

    return paragraphs


def remove_punctuation(s):
    """去除所有标点符号，保留文字"""
    return re.sub(r'[，。！？、；：""\'\'（）【】《》\[\]{}·—…\s]', '', s)


def find_paragraph_position(content, text, keywords=None):
    """
    在原文中查找段落位置
    策略：完全匹配 -> 关键词定位 -> 部分匹配
    返回: (start, end) 元组，未找到返回 (None, None)
    """
    if not text:
        return None, None

    # 策略1：完全匹配（带句号）
    if text in content:
        idx = content.find(text)
        return idx, idx + len(text)

    # 策略2：去掉句号匹配
    clean = text.rstrip('。')
    if clean in content:
        idx = content.find(clean)
        return idx, idx + len(clean)

    # 策略3：关键词定位（使用分析结果中的关键词）
    if keywords:
        for kw in keywords:
            if kw in content and kw in text:
                # 在原文中找到关键词位置
                kw_idx = content.find(kw)
                # 向前找句子开始
                start = kw_idx
                while start > 0 and content[start-1] not in '。\n':
                    start -= 1
                # 向后找句子结束（句号或换行）
                end = kw_idx
                while end < len(content) and content[end] != '。':
                    end += 1
                if end < len(content):
                    end += 1  # 包含句号
                return start, end

    # 策略4：部分匹配（文本开头部分在原文中）
    # 取前20个字符尝试匹配
    prefix = clean[:min(20, len(clean)//2)]
    if prefix and prefix in content:
        idx = content.find(prefix)
        # 从匹配位置向后扩展到完整句子
        end = idx
        while end < len(content) and content[end] not in '。\n':
            end += 1
        if end < len(content) and content[end] == '。':
            end += 1
        return idx, end

    return None, None


def get_marker(score, text):
    """根据分数返回高亮标记"""
    if not text or len(text.strip()) == 0:
        return ""
    if score >= 80:
        return f"=={text}=="
    elif score >= 60:
        return f"=={text}==*"
    else:
        return f"=={text}==**"


def highlight_content(content, paragraphs):
    """
    在内容中高亮相关段落
    策略：完全匹配 -> 关键词定位 -> 部分匹配
    返回: (highlighted_content, unmatched_list)
    """
    highlighted = content
    unmatched = []

    for para in paragraphs:
        original = para['text'].strip()
        score = para['score']
        keywords = para.get('keywords', [])

        if not original:
            continue

        start, end = find_paragraph_position(content, original, keywords)

        if start is not None:
            original_text = content[start:end]
            marker = get_marker(score, original_text)
            highlighted = highlighted[:start] + marker + highlighted[end:]
        else:
            # 记录未匹配段落
            unmatched.append({
                'text': original[:60] + '...' if len(original) > 60 else original,
                'score': score,
                'keywords': keywords
            })

    return highlighted, unmatched


def save_docx(content, title, output_path):
    """保存为Word文档"""
    if not DOCX_AVAILABLE:
        print(f"python-docx 未安装，无法保存: {output_path}")
        return False

    doc = Document()
    doc.add_heading(title, 0)

    for line in content.split('\n'):
        line = line.strip()
        if not line or line.startswith('#'):
            continue

        para = doc.add_paragraph()
        parse_and_add(para, line)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return True


def parse_and_add(para, text):
    """解析高亮标记并添加到段落"""
    pattern = r"==([^=]+)==(\*?\*)?"

    last_end = 0
    for match in re.finditer(pattern, text):
        before = text[last_end:match.start()]
        if before:
            para.add_run(before)

        highlighted = match.group(1)
        suffix = match.group(2) or ""

        run = para.add_run(highlighted)
        run.font.size = Pt(12)

        if not suffix:
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.font.bold = True
        elif suffix == '*':
            run.font.color.rgb = RGBColor(200, 150, 0)
        else:
            run.font.color.rgb = RGBColor(200, 0, 0)

        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        para.add_run(remaining)


def convert_analysis_to_docx(analysis_path, output_path):
    """将分析结果md文件转换为Word文档"""
    if not DOCX_AVAILABLE:
        print("python-docx 未安装，无法转换")
        return False

    with open(analysis_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    doc.add_heading('政策分析结果', 0)

    # 提取总分
    total_score_match = re.search(r'> \*\*总分\*\*[：:]\s*([\d.]+)', content)
    if total_score_match:
        score = total_score_match.group(1)
        p = doc.add_paragraph()
        run = p.add_run(f"总分: {score}/100")
        run.font.bold = True
        run.font.size = Pt(14)

    lines = content.split('\n')
    in_paragraph_section = False

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # 标题处理
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#'):
            doc.add_heading(line[1:], level=0)
        # 跳过总分行（已处理）
        elif '总分' in line and '**' in line:
            continue
        # 原文段落
        elif line.startswith('> ') and not line.startswith('> -'):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.add_run(text)
            in_paragraph_section = True
        # 评分行
        elif line.startswith('> -') and in_paragraph_section:
            score_match = re.search(r'评分[：:]\s*(\d+)', line)
            kw_match = re.search(r'关键词[：:]\s*(.+)', line)
            field_match = re.search(r'语义关联[：:]\s*(.+)', line)

            info_parts = []
            if score_match:
                info_parts.append(f"评分: {score_match.group(1)}/100")
            if field_match:
                info_parts.append(f"领域: {field_match.group(1)}")
            if kw_match:
                info_parts.append(f"关键词: {kw_match.group(1)}")

            if info_parts:
                p = doc.add_paragraph()
                p.add_run(' | '.join(info_parts)).italic = True
        # 摘要行
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.add_run(line[2:].strip())
        # 其他内容
        else:
            in_paragraph_section = False
            p = doc.add_paragraph()
            p.add_run(line)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return True


def highlight_doc(doc_path, verbose=True):
    """
    高亮单个政策文档
    返回: (success: bool, unmatched_count: int, unmatched_details: list)
    """
    if not os.path.exists(doc_path):
        print(f"文件不存在: {doc_path}")
        return False, 0, []

    doc_name = os.path.basename(doc_path)
    doc_title = os.path.splitext(doc_name)[0]

    # 查找分析结果
    analysis_path = find_analysis_result(doc_title)
    if not analysis_path:
        print(f"未找到分析结果: {doc_title}")
        return False, 0, []

    paragraphs = parse_paragraphs(analysis_path)
    if not paragraphs:
        print(f"无相关段落: {doc_title}")
        return False, 0, []

    # 读取原文
    with open(doc_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 高亮
    highlighted, unmatched = highlight_content(content, paragraphs)

    # 统计
    matched = len(paragraphs) - len(unmatched)
    green = len(re.findall(r'==[^=]+==(?!\*)', highlighted))
    yellow = len(re.findall(r'==[^=]+==\*', highlighted))
    red = len(re.findall(r'==[^=]+==\*\*', highlighted))

    if verbose:
        print(f"\n{'='*50}")
        print(f"文档: {doc_title}")
        print(f"解析段落: {len(paragraphs)} 个")
        print(f"匹配成功: {matched} 个")
        print(f"匹配失败: {len(unmatched)} 个")
        print(f"  - 绿色(>=80): {green} 处")
        print(f"  - 黄色(60-79): {yellow} 处")
        print(f"  - 红色(<60): {red} 处")

        if unmatched:
            print(f"\n未匹配段落 ({len(unmatched)} 个):")
            for i, u in enumerate(unmatched, 1):
                kw = ', '.join(u['keywords'][:3]) if u['keywords'] else '无'
                print(f"  {i}. [{u['score']}分] {u['text']}")
                print(f"     关键词: {kw}")

    # 保存
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, doc_name.replace('.md', '.docx'))

    if save_docx(highlighted, doc_title, output_path):
        print(f"\n已生成: {output_path}")

        # 提取总分用于文件名
        with open(analysis_path, 'r', encoding='utf-8') as f:
            analysis_content = f.read()
        score_match = re.search(r'> \*\*总分\*\*[：:]\s*([\d.]+)', analysis_content)
        score_suffix = f"_{score_match.group(1)}" if score_match else ""

        # 将分析结果md也转换成word，保存到analyze_result目录
        analysis_output_path = os.path.join(BASE_DIR, 'analyze_result', f"{doc_title}_分析结果{score_suffix}.docx")
        if convert_analysis_to_docx(analysis_path, analysis_output_path):
            print(f"已生成: {analysis_output_path}")
    else:
        md_path = os.path.join(OUTPUT_DIR, doc_name)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(highlighted)
        print(f"\n已生成(MD): {md_path}")

    return True, len(unmatched), unmatched


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python backend/highlight.py policy_document/xxx.md")
        sys.exit(1)

    doc_path = sys.argv[1]
    if not os.path.isabs(doc_path):
        doc_path = os.path.join(BASE_DIR, doc_path)

    success, unmatched_count, unmatched = highlight_doc(doc_path)
    sys.exit(0 if success else 1)
